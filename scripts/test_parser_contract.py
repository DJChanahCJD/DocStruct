from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from core.chunker import split_parse_result_into_chunks
from core.parser import DocxParser, ParserFactory


class ParserContractTest(unittest.TestCase):
    def test_docx_parser_preserves_block_order(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            file_path = Path(temp_dir) / "sample.docx"
            document = Document()
            document.add_heading("系统测试报告", level=1)
            document.add_paragraph("第一段说明")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "模块"
            table.cell(0, 1).text = "状态"
            table.cell(1, 0).text = "登录"
            table.cell(1, 1).text = "通过"
            document.add_paragraph("第二段结论")
            document.save(file_path)

            result = DocxParser().parse_to_result(str(file_path))

        block_types = [block.type for block in result.blocks]
        self.assertEqual(["title", "paragraph", "table", "paragraph"], block_types)
        self.assertIn("| 模块 | 状态 |", result.markdown)
        self.assertLess(result.markdown.find("第一段说明"), result.markdown.find("| 模块 | 状态 |"))
        self.assertLess(result.markdown.find("| 登录 | 通过 |"), result.markdown.find("第二段结论"))

    def test_parser_factory_parse_stays_backward_compatible(self) -> None:
        parser = ParserFactory.get_parser("static/examples/test_case_example.md")
        markdown = parser.parse("static/examples/test_case_example.md")
        self.assertIsInstance(markdown, str)
        self.assertIn("# 电商平台功能测试用例集", markdown)

    def test_parse_result_chunks_use_block_structure(self) -> None:
        parser = ParserFactory.get_parser("static/examples/test_case_example.md")
        parse_result = parser.parse_to_result("static/examples/test_case_example.md")
        chunks = split_parse_result_into_chunks(parse_result, max_chars=500, overlap_chars=50)
        self.assertGreater(len(chunks), 0)
        self.assertTrue(any(chunk.title_path for chunk in chunks))


if __name__ == "__main__":
    unittest.main()
